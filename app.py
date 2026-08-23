import streamlit as st
import markdown
import os
import re
import base64
import csv
import unicodedata
import textwrap
from io import StringIO, BytesIO
from docx import Document as DocxDocument
from pypdf import PdfReader
from groq import Groq
from langchain_groq import ChatGroq
from langchain_core.prompts import PromptTemplate
from dotenv import load_dotenv
from fpdf import FPDF

# Load environment variables
load_dotenv()

# --- Configuration & Initialization ---
API_KEY = os.environ.get("GROQ_API_KEY")
if not API_KEY:
    st.error("🚨 GROQ_API_KEY is missing! Please configure your .env file.")
    st.stop()

# Standard client used for Vision API calls
vision_client = Groq(api_key=API_KEY)

# --- Document Processing ---
class DocumentProcessor:
    @staticmethod
    def extract_from_pdf(file):
        file.seek(0)
        reader = PdfReader(file)
        pages_text = [page.extract_text() for page in reader.pages if page.extract_text()]
        return "\n\n".join(pages_text)

    @staticmethod
    def extract_from_docx(file):
        """Extracts text from Microsoft Word documents."""
        doc = DocxDocument(file)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

    @staticmethod
    def extract_from_text(file):
        """Instantly parses CSV and TXT files using UTF-8 decoding."""
        file.seek(0)
        return file.read().decode('utf-8', errors='ignore')

    @staticmethod
    def extract_from_image(file):
        file.seek(0)
        base64_image = base64.b64encode(file.read()).decode('utf-8')
        
        response = vision_client.chat.completions.create(
            model="qwen/qwen3.6-27b", 
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text", 
                            "text": (
                                "You are a strict OCR engine. Extract ONLY the visible text in this image. "
                                "Do NOT write introductions. Do NOT describe the visual content of the image. "
                                "If there is absolutely no readable text, you MUST output exactly and only: NO_TEXT_DETECTED"
                            )
                        },
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}},
                    ],
                }
            ],
            temperature=0.0
        )
        
        result = (response.choices[0].message.content or "").strip()
        lower_res = result.lower()
        
        if "no_text_detected" in lower_res or "###none###" in lower_res or lower_res == "empty":
            return ""
            
        apologies = [
            "no text", "absence of text", "cannot see any text", 
            "no readable text", "does not contain any text", "no alphanumeric",
            "contains no readable text", "nothing to transcribe", "no visible text"
        ]
        
        if any(phrase in lower_res for phrase in apologies):
            return ""

        prefixes = ["the text in the image is:", "visible text:", "extracted text:", "the image contains:"]
        for prefix in prefixes:
            if lower_res.startswith(prefix):
                result = result[len(prefix):].strip()
                
        return result

# --- LLM Orchestration ---
class SummarizerEngine:
    @staticmethod
    def generate(text, length):
        words = text.split()
        
        # --- AGGRESSIVE SHORT-TEXT BYPASS ---
        if len(words) < 30:
            return (
                f"**Extracted Content:**\n> {text}\n\n"
                "*Note: This document is too brief to generate a structured, multi-paragraph summary or improvement suggestions.*"
            )
            
        # --- CONTEXT WINDOW PROTECTOR ---
        if len(text) > 10000:
            text = text[:10000]
            st.warning("⚠️ **Context Window Optimization:** This document exceeds standard compute thresholds. To maintain high fidelity and optimize response latency, the data has been intelligently truncated to analyze the primary sections.")
            
        llm = ChatGroq(temperature=0.2, model="openai/gpt-oss-120b")
        
        # Controls the paragraph length
        length_guidelines = {
            "Short": "CRITICAL LENGTH LIMIT: You MUST write exactly 2 to 3 sentences. Absolute maximum of 60 words. Be extremely brief, direct, and punchy.",
            "Medium": "LENGTH LIMIT: Write exactly 2 paragraphs (approx 150-200 words).",
            "Long": "LENGTH LIMIT: Write a deeply comprehensive, highly detailed multi-paragraph analysis (400+ words)."
        }

        # Controls the table row limits
        data_guidelines = {
            "Short": "DATA LIMIT: Extract ONLY the top 3 most critical rows or items into a single, tiny Markdown table. Ignore all minor details and secondary tables.",
            "Medium": "DATA LIMIT: Extract the most relevant data points into a concise Markdown table (up to 5-8 key rows).",
            "Long": "DATA LIMIT: Extract all available schedules, qualifications, or tabular data exhaustively into comprehensive Markdown tables."
        }
        
        # THE ULTIMATE PROMPT
        template = """
        You are an Elite Executive Analyst and Document Intelligence Expert. Your task is to provide a comprehensive, highly accurate analysis of the provided text, which may come from a PDF, Word Document, CSV dataset, Plain Text file, or OCR-extracted Image.

        CRITICAL ADAPTATION RULES:
        1. For Reports/Resumes/Prose (PDF/DOCX): Synthesize the core value proposition, strategic importance, and primary narrative themes.
        2. For Spreadsheets/Datasets (CSV): Analyze the data structure, identify statistical trends, extract key metrics, and explain what the dataset represents.
        3. For Images/Scanned Documents (OCR Text): Intelligently auto-correct any OCR scanning artifacts (e.g., misspelled words, broken lines) and summarize the core visible information.
        4. For Code/System Logs (TXT): Explain the functionality, identify key errors, or summarize the system state.

        Please adhere STRICTLY to the following structure. Do not include any conversational filler.

        ### Executive Summary
        ({length_instruction})
        (Focus on the essence, primary trends, or core purpose while avoiding mere repetition.)

        ### Key Highlights
        - (Insightful bullet point 1: Summarize a major achievement, core theme, critical data trend, or primary image subject.)
        - (Insightful bullet point 2: Highlight another significant aspect, finding, or data anomaly.)
        - (Insightful bullet point 3: Present a third key point encapsulating important information.)

        ### Extracted Structured Data
        ({data_instruction})
        (If the original document IS a CSV/spreadsheet, format the rows/columns based on the limit above. If it is an image of a receipt/invoice, extract line items based on the limit above. If absolutely no structured data exists, state "No structured data found.")

        ### 💡 Improvement Suggestions
        1. (Provide a specific recommendation to enhance the document's impact, OR for datasets/logs/images, suggest ways to improve data quality, formatting, or error handling.)
        2. (Provide a second specific and actionable recommendation.)

        ---
        DOCUMENT TEXT TO SUMMARIZE:
        {text}
        """
        
        # Added data_instruction to the input variables here!
        prompt = PromptTemplate(input_variables=["length_instruction", "data_instruction", "text"], template=template)
        chain = prompt | llm
        
        # Pass both constraints into the chain
        return chain.invoke({
            "length_instruction": length_guidelines[length],
            "data_instruction": data_guidelines[length],
            "text": text
        }).content
    
# --- Text Sanitizer ---
def sanitize_llm_text(text):
    """Cleans up tricky LLM unicode characters and HTML tags before parsing."""
    replacements = {
        '\u2018': "'", '\u2019': "'", '\u201c': '"', '\u201d': '"',
        '\u2013': '-', '\u2014': '-', '\u2026': '...', '\u00A0': ' ',
        '\u2011': '-', '\u202f': ' ', '\u200b': '', '💡': 'Tip:',
        '<br>': '\n', '<br/>': '\n', '<br />': '\n'
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
        
    return unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('ascii')

# --- PDF Generator ---
def generate_pdf_summary(text_input):
    """Converts the LLM Markdown summary into a safe, natively formatted PDF."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    safe_text = sanitize_llm_text(text_input)
    table_data = []
    is_in_table = False
    
    # Prevents the "Dangling Cursor" bug
    def print_safe(pdf_obj, txt, style="", size=11):
        # 1. Hard-reset the cursor to the left margin before EVERY print!
        pdf_obj.set_x(pdf_obj.l_margin)
        pdf_obj.set_font("Helvetica", style, size)
        
        # 2. Chop massive unbreakable words (like URLs or non-breaking hyphens)
        safe_txt = re.sub(r'([^\s]{35})', r'\1 ', txt)
        
        try:
            pdf_obj.multi_cell(0, 6, text=safe_txt)
        except Exception:
            # 3. If multi_cell still crashes, fallback to write() which handles line-breaks naturally
            pdf_obj.set_x(pdf_obj.l_margin)
            try:
                pdf_obj.write(6, text=safe_txt + '\n')
            except Exception:
                pass
                
    for line in safe_text.split('\n'):
        line = line.strip()
        
        if is_in_table and not line.startswith('|'):
            pdf.set_font("Helvetica", size=9)
            try:
                with pdf.table(text_align="LEFT") as table:
                    for row_cells in table_data:
                        row = table.row()
                        for cell_text in row_cells:
                            safe_cell = re.sub(r'([^\s]{20})', r'\1 ', cell_text)
                            row.cell(safe_cell)
            except Exception:
                pass
            table_data = []
            is_in_table = False
            pdf.ln(5)

        if not line:
            continue
            
        if len(line) >= 3 and all(c in '-_*= \t' for c in line):
            pdf.ln(4)
            continue
            
        if line.startswith('#'):
            clean_header = line.lstrip('# ').replace('**', '')
            print_safe(pdf, clean_header, "B", 14)
            pdf.ln(2)
            
        elif line.startswith('|'):
            if '---' in line or '===' in line:
                continue
            cells = [c.strip().replace('**', '') for c in line.split('|')]
            if cells and cells[0] == '': cells.pop(0)
            if cells and cells[-1] == '': cells.pop()
            if cells:
                is_in_table = True
                table_data.append(cells)
                
        elif line.startswith('- ') or line.startswith('* '):
            print_safe(pdf, "  - " + line.replace('**', '').lstrip('- *'))
            
        else:
            print_safe(pdf, line.replace('**', ''))
            
    if is_in_table and table_data:
        pdf.set_font("Helvetica", size=9)
        try:
            with pdf.table(text_align="LEFT") as table:
                for row_cells in table_data:
                    row = table.row()
                    for cell_text in row_cells:
                        safe_cell = re.sub(r'([^\s]{20})', r'\1 ', cell_text)
                        row.cell(safe_cell)
        except Exception:
            pass

    return bytes(pdf.output())

def generate_docx_summary(text_input):
    """Converts the Markdown summary into a native Microsoft Word document with real tables."""
    safe_text = sanitize_llm_text(text_input)
        
    doc = DocxDocument()
    doc.add_heading('Executive Summary', level=1)
    
    table = None 
    
    for line in safe_text.split('\n'):
        line = line.strip()
        if not line:
            table = None 
            continue
            
        # Catch ANY size header for Word documents too
        if line.startswith('#'):
            table = None
            clean_header = line.lstrip('# ').replace('**', '')
            doc.add_heading(clean_header, level=2)
            
        elif line.startswith('|'):
            if '---' in line or '===' in line: 
                continue 
                
            cells = [c.strip().replace('**', '') for c in line.split('|')]
            if cells and cells[0] == '': cells.pop(0)
            if cells and cells[-1] == '': cells.pop()
            
            if not cells:
                continue
                
            if not table:
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, cell_text in enumerate(cells):
                    if i < len(hdr_cells):
                        hdr_cells[i].text = cell_text
            else:
                while len(cells) < len(table.columns): 
                    cells.append("")
                    
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(cells[:len(table.columns)]):
                    row_cells[i].text = cell_text
                        
        elif line.startswith('- ') or line.startswith('* '):
            table = None
            clean_bullet = line.replace('**', '').lstrip('- *')
            doc.add_paragraph(clean_bullet, style='List Bullet')
            
        else:
            table = None
            doc.add_paragraph(line.replace('**', ''))
            
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def generate_csv_summary(text_input):
    """Converts the summary into a structured CSV spreadsheet."""
    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(["Section", "Content"]) 
    
    current_section = "General Summary"
    for line in text_input.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('###'):
            current_section = line.replace('###', '').replace('**', '').strip()
        else:
            clean_line = line.replace('**', '').replace('- ', '').replace('* ', '')
            writer.writerow([current_section, clean_line])
            
    return output.getvalue()

def main():
    # Forces the sidebar to be open by default
    st.set_page_config(page_title="Document Summary Assistant", page_icon="⚡", layout="centered", initial_sidebar_state="expanded")
    
    st.markdown("""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');
        html, body, [class*="css"] { font-family: 'Inter', sans-serif !important; }
        
        /* Hide Defaults */
        #MainMenu, footer, .stDeployButton { display: none !important; }
        
        /* App Background */
        .stApp {
            background-color: #09090b !important;
            background-image: 
                radial-gradient(at 0% 0%, rgba(99, 102, 241, 0.15) 0px, transparent 50%),
                radial-gradient(at 100% 0%, rgba(168, 85, 247, 0.15) 0px, transparent 50%);
            background-attachment: fixed;
            color: #fafafa !important;
        }
        
        /* --- BRAND NEW UI UPGRADES --- */
        /* 1. Make the top header transparent to remove the black bar */
        header { background: transparent !important; }

        /* 2. Style the ">>" arrow to look like a purple UI button */
        [data-testid="collapsedControl"] {
            color: #c084fc !important;
            background-color: rgba(24, 24, 27, 0.6) !important;
            border: 1px solid rgba(168, 85, 247, 0.4) !important;
            border-radius: 8px !important;
            padding: 0px 12px !important;
            margin: 15px !important;
            display: flex !important;
            align-items: center !important;
            transition: all 0.3s ease !important;
        }
        
        /* 3. Inject a label next to the ">>" arrow so users know what it is */
        [data-testid="collapsedControl"]::after {
            content: "Summary Length";
            font-family: 'Inter', sans-serif;
            font-weight: 600;
            font-size: 0.9rem;
            margin-left: 8px;
        }
        
        /* Hover effect for the new button */
        [data-testid="collapsedControl"]:hover {
            background-color: rgba(168, 85, 247, 0.2) !important;
            border-color: #c084fc !important;
        }
        /* ----------------------------- */

        /* Container */
        .block-container { max-width: 800px !important; padding-top: 3rem !important; }
        
        /* Gradient Button */
        div.stButton > button:first-child {
            background: linear-gradient(90deg, #6366f1, #a855f7) !important;
            color: white !important;
            border: none !important;
            border-radius: 8px !important;
            padding: 0.75rem !important;
            font-weight: 600 !important;
            transition: all 0.3s ease !important;
            box-shadow: 0 4px 15px rgba(168, 85, 247, 0.3) !important;
        }
        div.stButton > button:first-child:hover {
            transform: translateY(-2px) !important;
            box-shadow: 0 8px 25px rgba(168, 85, 247, 0.5) !important;
        }
        
        /* Dropzone & Uploaded File Card */
        [data-testid="stFileUploadDropzone"] {
            background: rgba(24, 24, 27, 0.5) !important;
            border: 1px dashed rgba(168, 85, 247, 0.5) !important;
            border-radius: 12px !important;
            backdrop-filter: blur(8px);
        }
        [data-testid="stUploadedFile"] {
            background: rgba(168, 85, 247, 0.1) !important;
            border: 1px solid rgba(168, 85, 247, 0.4) !important;
            border-radius: 8px !important;
        }
        
        /* Sidebar & Expander Details */
        [data-testid="stSidebar"] { background-color: rgba(9, 9, 11, 0.95) !important; border-right: 1px solid rgba(255,255,255,0.05) !important; }
        [data-testid="stExpander"] details {
            background: rgba(24, 24, 27, 0.6) !important;
            border: 1px solid rgba(255,255,255,0.1) !important;
            border-radius: 8px !important;
        }
        [data-testid="stExpander"] details summary { color: #c084fc !important; font-weight: 600 !important; }
        
        /* Custom Alert Components */
        .glass-info {
            background: rgba(139, 92, 246, 0.15);
            border: 1px solid rgba(139, 92, 246, 0.3);
            border-radius: 8px; padding: 12px; color: #e9d5ff; font-size: 0.9rem;
            backdrop-filter: blur(10px);
        }
        .glass-success {
            background: rgba(16, 185, 129, 0.15);
            border: 1px solid rgba(16, 185, 129, 0.3);
            border-radius: 8px; padding: 14px; color: #a7f3d0; font-weight: 600;
            display: flex; align-items: center; gap: 8px; margin-bottom: 1rem;
            backdrop-filter: blur(10px);
        }

        /* SaaS Data Table Styling */
        .stMarkdown table {
            width: 100% !important;
            border-collapse: collapse !important;
            margin: 1.5rem 0 !important;
            border-radius: 8px !important;
            overflow: hidden !important;
            border: 1px solid rgba(255, 255, 255, 0.1) !important;
        }
        .stMarkdown th {
            background-color: rgba(139, 92, 246, 0.15) !important; 
            color: #e9d5ff !important;
            font-weight: 600 !important;
            padding: 12px 16px !important;
            text-align: left !important;
            border-bottom: 1px solid rgba(255, 255, 255, 0.1) !important;
        }
        .stMarkdown td {
            padding: 12px 16px !important;
            border-bottom: 1px solid rgba(255, 255, 255, 0.05) !important;
            background-color: rgba(24, 24, 27, 0.5) !important; 
            color: #e4e4e7 !important;
            transition: background-color 0.2s ease !important;
        }
        .stMarkdown tr:last-child td { border-bottom: none !important; }
        .stMarkdown tr:hover td { background-color: rgba(139, 92, 246, 0.08) !important; }
        </style>
    """, unsafe_allow_html=True)
    
    st.markdown("""
        <h1 style='font-weight: 800; font-size: 3rem; background: linear-gradient(to right, #fff, #a1a1aa); -webkit-background-clip: text; -webkit-text-fill-color: transparent; margin-bottom: 0px;'>Document Summary Assistant</h1>
        <p style='color: #a1a1aa; font-size: 1.1rem; margin-bottom: 2rem;'>Upload any document (PDF, DOCX, CSV, TXT, or Image) up to 25MB to generate intelligent, structured summaries powered by Groq.</p>
    """, unsafe_allow_html=True)
    
    with st.sidebar:
        # Changed symbol and text, removed the 'glass-info' section entirely!
        st.markdown("<h2 style='color: white; font-size: 1.2rem;'>📏 Summary Settings</h2>", unsafe_allow_html=True)
        summary_length = st.radio("Select Summary Length:", ("Short", "Medium", "Long")) 
        st.markdown("<br>", unsafe_allow_html=True)

    uploaded_files = st.file_uploader(
        "Drop your document(s) here", 
        type=["pdf", "png", "jpg", "jpeg", "csv", "txt", "docx"], 
        accept_multiple_files=True, 
        label_visibility="collapsed"
    )

    # (The rest of your code like 'if "final_summary" not in st.session_state:' stays exactly the same)

    if "final_summary" not in st.session_state:
        st.session_state.final_summary = None

    if uploaded_files and st.button("Generate Smart Summary", type="primary", use_container_width=True):
        
        st.session_state.final_summary = None
        
        with st.spinner(f"Processing {len(uploaded_files)} document(s)..."):
            try:
                combined_raw_text = ""
                
                # Loop through all uploaded files
                for uploaded_file in uploaded_files:
                    file_ext = uploaded_file.name.lower().split('.')[-1]
                    
                    if file_ext == "pdf":
                        raw_text = DocumentProcessor.extract_from_pdf(uploaded_file)
                    elif file_ext == "docx":
                        raw_text = DocumentProcessor.extract_from_docx(uploaded_file)
                    elif file_ext in ["csv", "txt"]:
                        raw_text = DocumentProcessor.extract_from_text(uploaded_file)
                    else:
                        raw_text = DocumentProcessor.extract_from_image(uploaded_file)
                    
                    if raw_text and raw_text.strip():
                        combined_raw_text += f"\n\n--- Document: {uploaded_file.name} ---\n{raw_text}\n"
                
                if not combined_raw_text.strip():
                    st.error("⚠️ **No readable text detected.** Please upload valid documents containing text.")
                    return

                with st.expander("🔍 View Extracted Raw Text"):
                    st.text(combined_raw_text)
                    
                summary = SummarizerEngine.generate(combined_raw_text, summary_length)
                st.session_state.final_summary = summary
                
            except Exception as e:
                error_msg = str(e)
                if "413" in error_msg or "rate_limit_exceeded" in error_msg:
                    st.error("🚨 **System Overload:** The document payload is too dense for the current compute tier. Please try fewer or smaller files.")
                else:
                    st.error(f"🚨 **Processing Error:** An unexpected issue occurred during analysis. Details: {error_msg}")

    # Display from State & Add Download Button
    if st.session_state.final_summary:
        st.markdown('<div class="glass-success"><span>✨</span> Analysis Complete!</div>', unsafe_allow_html=True)
        st.markdown("### 📋 Executive Summary")
        st.markdown(st.session_state.final_summary)
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        # --- THE MULTI-FORMAT EXPORT HUB ---
        col1, col2 = st.columns([1, 2])
        
        with col1:
            export_format = st.selectbox(
                "Select Export Format:",
                options=[
                    "Word Document (.docx)", 
                    "PDF (.pdf)", 
                    "CSV Spreadsheet (.csv)", 
                    "Markdown (.md)", 
                    "Plain Text (.txt)"
                ],
                label_visibility="collapsed"
            )
            
        with col2:
            if export_format == "Word Document (.docx)":
                docx_bytes = generate_docx_summary(st.session_state.final_summary)
                st.download_button(
                    label="📥 Download Summary",
                    data=docx_bytes,
                    file_name="executive_summary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            elif export_format == "PDF (.pdf)":
                pdf_bytes = generate_pdf_summary(st.session_state.final_summary)
                st.download_button(
                    label="📥 Download Summary",
                    data=pdf_bytes,
                    file_name="executive_summary.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            elif export_format == "CSV Spreadsheet (.csv)":
                csv_data = generate_csv_summary(st.session_state.final_summary)
                st.download_button(
                    label="📥 Download Summary",
                    data=csv_data,
                    file_name="executive_summary.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            elif export_format == "Markdown (.md)":
                st.download_button(
                    label="📥 Download Summary",
                    data=str(st.session_state.final_summary),
                    file_name="executive_summary.md",
                    mime="text/markdown",
                    use_container_width=True
                )
            else:
                summary = st.session_state.final_summary
                txt_data = (
                    summary.replace("**", "").replace("###", "")
                    if isinstance(summary, str)
                    else "\n".join(str(item) for item in summary)
                )
                st.download_button(
                    label="📥 Download Summary",
                    data=txt_data,
                    file_name="executive_summary.txt",
                    mime="text/plain",
                    use_container_width=True
                )

if __name__ == "__main__":
    main()